try:
    from docx import Document
    from docx.shared import Inches
except ImportError:
    print("DOCX conversion Unavailable")
try:
    from bs4 import BeautifulSoup
except ImportError:
    print("HTML Conversion Unavailable")
try:
    import markdown
except ImportError:
    print("HTML Conversion Unavailable")
    print("Markdown Conversion Unavailable")
import re
import collections
import os
try:
    from reportlab.pdfgen import canvas
    from reportlab.lib.units import cm
    from reportlab.lib.pagesizes import A4
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Image, tableofcontents
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
except ImportError:
    print("PDF Generation Unavailable")

def strip_html(stringIn):
    stringIn = str(stringIn)
    stripped_item = re.sub('<[^<]+?>', '', stringIn)
    return stripped_item

def html_to_json(htmlString,outFile=None):
    """Convert HTML to pywriter's JSON syntax"""
    htmlString = markdown.markdown(htmlString)
    htmlString = BeautifulSoup(htmlString, 'html.parser')
    finalDict = {}
    finalDict['document'] = collections.OrderedDict()
    for item in htmlString:
        stringItem = str(item)
        if '<h1' in stringItem:
            stripped_item = strip_html(item)
            finalDict['document'][stripped_item] = 'title'
        elif '<h2' in stringItem:
            stripped_item = strip_html(item)
            finalDict['document'][stripped_item] = 'subtitle'
        elif '<h3' in stringItem:
            stripped_item = strip_html(item)
            finalDict['document'][stripped_item] = 'tagline'
        elif '<b' in stringItem:
            stripped_item = strip_html(item)
            finalDict['document'][stripped_item] = 'bold'
        elif '<strong' in stringItem:
            stripped_item = strip_html(item)
            finalDict['document'][stripped_item] = 'bold'
        elif '<i' in stringItem:
            stripped_item = strip_html(item)
            finalDict['document'][stripped_item] = 'italic'
        elif '<em' in stringItem:
            stripped_item = strip_html(item)
            finalDict['document'][stripped_item] = 'italic'
        elif '<li' in stringItem:
            stripped_item = strip_html(item)
            finalDict['document'][stripped_item] = 'bulletpoint'
        elif '<p' in stringItem:
            stripped_item = strip_html(item)
            finalDict['document'][stripped_item] = 'paragraph'
        if outFile:
            finalDict['outFile'] = outFile
        else:
            folder = os.getcwd()
            if '\\' in folder:
                folder = folder.replace('\\','/')
            finalDict['outFile'] = folder + '/save.docx'
    return finalDict

def generate_docx(dictIn):
    """Takes either a list for unformatted text, or a dict with formatting options"""
    try:
        docStruct = dictIn['document']
    except KeyError:
        docStruct = list(dictIn)
    documentReal = Document()

    # Try and build the document from a dict
    try:
        previousPara = False
        for key, value in docStruct.items():
            if value == 'paragraph':
                if not previousPara:
                    paragraph = documentReal.add_paragraph('')
                paragraph.add_run(key)
                previousPara = True
            elif value == 'title':
                previousPara = False
                documentReal.add_heading(key, level=0)
            elif value == 'subtitle':
                previousPara = False
                documentReal.add_heading(key, level=1)
            elif value == 'tagline':
                previousPara = False
                documentReal.add_heading(key, level=2)
            elif value == 'bold':
                if not previousPara:
                    paragraph = documentReal.add_paragraph('')
                paragraph.add_run(key).bold = True
            elif value == 'italic':
                if not previousPara:
                    paragraph = documentReal.add_paragraph('')
                paragraph.add_run(key).italic = True
            elif value == 'bulletpoint':
                paragraph = documentReal.add_paragraph('')
                prun = paragraph.add_run('* ' + key)
            elif value == 'image':
                documentReal.add_picture(key, width=Inches(1.0))
            elif value == 'paragraph':
                if not previousPara:
                    documentReal.add_paragraph('')
                paragraph.add_run(key)

    # If not a dict, process it as a list
    except AttributeError:
        for item in docStruct:
            documentReal.add_paragraph(item)
    # Save the file from the dict value
    try:
        documentReal.save(dictIn['outFile'])
    except KeyError:
        # Or save into the current working directory.
        folder = os.getcwd()
        if '\\' in folder:
            folder = folder.replace('\\', '/')
        documentReal.save(folder + '/save.docx')

def generate_markdown(dictIn):
    """Takes either a list for unformatted text, or a dict with formatting options"""
    try:
        docStruct = dictIn['document']
    except AttributeError:
        docStruct = list(dictIn)

def addPageNumber(canvas, doc):
    page_num = canvas.getPageNumber()
    text = "Page %s" % page_num
    canvas.drawRightString(2*cm, 0.2*cm, text)

def generate_pdf(dictIn):
    """Takes either a list for unformatted text, or a dict with formatting options"""
    try:
        docStruct = dictIn['document']
    except TypeError:
        docStruct = list(dictIn)

    try:
        outFile = dictIn['outFile']
    except TypeError:
        folder = os.getcwd()
        if '\\' in folder:
            folder = folder.replace('\\','/')
        outFile = folder + '/save.pdf'

    doc = SimpleDocTemplate(outFile,pagesize=A4,rightMargin=72,leftMargin=72,topMargin=72,bottomMargin=18)
    styles=getSampleStyleSheet()
    documentReal = []
    try:
        for key, value in docStruct.items():
           if value == 'title':
               documentReal.append(Paragraph('<font size=21>%s</font>' % key, styles["Title"]))
           elif value == 'subtitle':
               documentReal.append(Paragraph('<font size=18>%s</font>' % key, styles["Heading1"]))
           elif value == 'tagline':
               documentReal.append(Paragraph('<font size=16>%s</font>' % key, styles["Heading2"]))
           elif value == 'bold':
               documentReal.append(Paragraph('<font size=16><b>%s</b></font>' % key, styles["Normal"]))
           documentReal.append(Spacer(1, 12))
        doc.build(documentReal, onFirstPage=addPageNumber, onLaterPages=addPageNumber)
    except AttributeError:
        for item in docStruct:
            documentReal.append(Paragraph('<font size=12>%s</font>' % item, styles["Normal"]))
            documentReal.append(Spacer(1, 12))
        doc.build(documentReal, onFirstPage=addPageNumber, onLaterPages=addPageNumber)
